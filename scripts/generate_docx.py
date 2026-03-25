#!/usr/bin/env python3
"""
Generate DOCX workbook with proper Chinese fonts.
Usage: python generate_docx_fixed.py --input workbook.md --output filename.docx
"""

import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='仿宋', font_size=12, bold=False, color=None):
    """Set Chinese font for a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def parse_and_create_docx(md_content, output_path):
    """Parse markdown and create DOCX with proper fonts."""
    doc = Document()
    
    # Set default document font
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(12)
    
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Title (九下...aespa Edition)
        if 'aespa Edition' in line and '九下' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '黑体', 16, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Date line
        elif line.startswith('整理日期：'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Part headers (# 第一部分/第二部分)
        elif line.startswith('# 第一部分') or line.startswith('# 第二部分'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('# ', ''))
            set_chinese_font(run, '黑体', 14, bold=True)
            
        # Sub headers (## 5. 答案与解析)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('## ', ''))
            set_chinese_font(run, '黑体', 12, bold=True)
            
        # Image markers
        elif line.startswith('📷'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 10, color=(128, 128, 128))
            
        # Member quotes with colors
        elif line.startswith('💙'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11, color=(0, 102, 204))
        elif line.startswith('🖤'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11, color=(102, 0, 153))
        elif line.startswith('❄️'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11, color=(0, 153, 204))
        elif line.startswith('🦋'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11, color=(204, 51, 0))
            
        # Signatures and ending
        elif 'aespa 签名' in line or 'We are aespa' in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 11, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Bold text (正确答案)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            text = line.strip('*')
            run = p.add_run(text)
            set_chinese_font(run, '仿宋', 12, bold=True, color=(0, 102, 204))
            
        # Horizontal rules (---)
        elif line == '---':
            doc.add_paragraph('─' * 40)
            
        # Regular text
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '仿宋', 12)
    
    doc.save(output_path)
    print(f"✓ DOCX saved with proper fonts: {output_path}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()
    
    with open(args.input, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    parse_and_create_docx(md_content, args.output)
